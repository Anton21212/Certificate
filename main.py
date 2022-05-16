import docx
import datetime
import time

mounth = {1: "Января",
          2: "Февраля",
          3: "Марта",
          4: "Апреля",
          5: "Мая",
          6: "Июня",
          7: "Июля",
          8: "Августа",
          9: "Сентрября",
          10: "Октярбя",
          11: "Ноября",
          12: "Декабря"}


class DOC:

    def __init__(self, document):
        self.document = document
        self.table = document.tables[0]

    def write_fio(self, input_name):
        par_1 = self.table.rows[0].cells[1].paragraphs[2]
        par_1.runs[0].text = input_name
        self.document.save("Sertifikat_1.docx")

    def write_date_start(self, date_start):
        start = datetime.datetime.strptime(date_start, "%Y/%m/%d")
        start = f" {start.day} {mounth[start.month]} {start.year} "
        return start

    def check_write_date_start(self, date_start):
        try:
            time.strptime(date_start, '%Y/%m/%d')
        except ValueError:
            print('Invalid date!')
        else:
            par_4 = self.table.cell(0, 1).paragraphs[3]
            par_4.runs[1].text = self.write_date_start(date_start)

    def check_write_date_finish(self, date_finish):
        try:
            time.strptime(date_finish, '%Y/%m/%d')
        except ValueError:
            print('Invalid date!')
        else:
            par_4 = self.table.cell(0, 1).paragraphs[3]
            par_4.runs[15].text = self.write_date_finish(date_finish)

    def write_date_finish(self, date_finish):
        finish = datetime.datetime.strptime(date_finish, "%Y/%m/%d")
        finish = f"{finish.day} {mounth[finish.month]} {finish.year}"
        return finish

    def write_date(self, date_start, date_finish):
        par_3 = self.table.rows[0].cells[1].paragraphs[3]

        for run in par_3.runs[1:12]:
            run.text = ''
        for run in par_3.runs[14:]:
            run.text = ''

        self.check_write_date_start(date_start)
        self.check_write_date_finish(date_finish)
        self.document.save("Sertifikat_1.docx")

    def write_time(self, t):
        time = self.table.rows[0].cells[1].paragraphs[6]
        time.runs[2].text = str(t)
        self.document.save("Sertifikat_1.docx")

    def check_write_time(self, t):
        if str(t).isnumeric():
            self.write_time(t)
        else:
            print("Введите тип инт")

    def write_programm(self, programm_input):
        programm = self.table.rows[0].cells[1].paragraphs[7]
        programm.runs[4].text = programm_input
        self.document.save("Sertifikat_1.docx")

    def run(self):
        self.write_fio(input_name=input("Введите ФИО: "))
        self.write_date(date_start=input("Введите дату начала: "), date_finish=input("Введите дату конца: "))
        self.check_write_time(t=input("Введите количество времени обучения: "))
        self.write_programm(programm_input=input("Введите название программы: "))


if __name__ == '__main__':
    DOC(docx.Document("Sertifikat.docx")).run()
